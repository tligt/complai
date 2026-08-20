import os
import io
import json
import uuid
import subprocess
import tempfile
import requests
import streamlit as st
from datetime import date
from rag import retrieve, Chunk
from database import get_supabase

# ── Constants ─────────────────────────────────────────────────

# Document types come from the obligation registry — a private copy here
# was one of the three competing definitions S23 removed.
from obligations import DOCUMENT_TYPES   # noqa: E402

LEGAL_FORMS = {
    "BE": ["SRL", "SA", "SNC", "SCS", "SC", "ASBL", "Fondation", "Other"],
    "FR": ["SARL", "SAS", "SASU", "SA", "SNC", "Auto-entrepreneur", "Association", "Other"],
    "NL": ["BV", "NV", "VOF", "Eenmanszaak", "Other"],
    "DE": ["GmbH", "AG", "UG", "OHG", "Other"],
    "LU": ["SARL", "SA", "SNC", "Other"],
    "EU": ["SRL", "SARL", "GmbH", "Ltd", "Other"],
}

DPA_CONTACTS = {
    "BE": "Autorité de protection des données (APD) — apd-gba.be",
    "FR": "Commission Nationale de l'Informatique et des Libertés (CNIL) — cnil.fr",
    "NL": "Autoriteit Persoonsgegevens (AP) — autoriteitpersoonsgegevens.nl",
    "DE": "Bundesdatenschutzbeauftragter (BfDI) — bfdi.bund.de",
    "LU": "Commission Nationale pour la Protection des Données (CNPD) — cnpd.public.lu",
    "EU": "The relevant national data protection authority in your country of establishment",
}


# ── Supabase helpers ──────────────────────────────────────────

def load_intake(client_id: str, user_id: str, document_type: str) -> dict:
    """Load previously saved intake data for this client + document type."""
    try:
        supabase = get_supabase()
        res = supabase.table("document_intake") \
            .select("*") \
            .eq("client_id", client_id) \
            .eq("user_id", user_id) \
            .eq("document_type", document_type) \
            .execute()
        return res.data[0] if res.data else {}
    except Exception:
        return {}


def save_intake(client_id: str, user_id: str, document_type: str, fields: dict) -> bool:
    """Save intake data — upsert on client_id + document_type."""
    try:
        supabase = get_supabase()
        record = {
            "user_id": user_id,
            "client_id": client_id,
            "document_type": document_type,
            **fields,
        }
        supabase.table("document_intake").upsert(
            record,
            on_conflict="client_id,document_type"
        ).execute()
        return True
    except Exception as e:
        st.error(f"Could not save intake data: {e}")
        return False


def update_client_profile(client_id: str, user_id: str, fields: dict) -> bool:
    """Update universal profile fields on the clients table."""
    try:
        supabase = get_supabase()
        supabase.table("clients") \
            .update(fields) \
            .eq("id", client_id) \
            .eq("user_id", user_id) \
            .execute()
        return True
    except Exception as e:
        st.error(f"Could not update client profile: {e}")
        return False


def save_document_record(user_id: str, client_id: str | None,
                          document_type: str, language: str,
                          company_name: str) -> bool:
    """Save a document generation record."""
    try:
        supabase = get_supabase()
        supabase.table("documents").insert({
            "user_id": user_id,
            "client_id": client_id,
            "document_type": document_type,
            "language": language,
            "company_name": company_name,
        }).execute()
        return True
    except Exception as e:
        return False


def load_document_history(user_id: str, client_id: str | None) -> list:
    """Load document generation history for a client."""
    try:
        supabase = get_supabase()
        q = supabase.table("documents") \
            .select("*") \
            .eq("user_id", user_id) \
            .order("generated_at", desc=True) \
            .limit(20)
        if client_id:
            q = q.eq("client_id", client_id)
        return q.execute().data or []
    except Exception:
        return []




# ── AI suggestion engine ──────────────────────────────────────

def save_document_with_files(
    user_id: str,
    client_id: str | None,
    document_type: str,
    language: str,
    company_name: str,
    docx_bytes: bytes,
    pdf_bytes: bytes | None = None,
    odt_bytes: bytes | None = None,
    *,
    xlsx_bytes: bytes | None = None,
    template_version_id: str | None = None,
    document_group_id: str | None = None,
    outstanding_fields: list | None = None,
    jurisdictions_applied: list | None = None,
    brand_profile_version: str | None = None,
) -> str | None:
    """Save document record and upload files to Supabase Storage. Returns document ID.

    S25 stamping (all optional, keyword-only):

      template_version_id   which template version produced this. Cheap to
                            record now, impossible to reconstruct later — a
                            document generated today is otherwise
                            indistinguishable from one generated against a
                            later revision.
      document_group_id     ties the language siblings of one generation
                            together. Two rows for a Brussels client are the
                            SAME document in two languages, not two documents,
                            and S27 adoption applies to the group.
      outstanding_fields    unresolved placeholders, per language. A French
                            body can carry one its Dutch sibling does not, so
                            completeness must be summed across the group.
      jurisdictions_applied which markets' rules were resolved into this text.
      brand_profile_version unused until S43; taken now so the call sites do
                            not need touching again.

    S26:
      xlsx_bytes            a spreadsheet representation of the SAME record.
                            Registers are tables and get filtered and sorted;
                            the CNIL publishes its model register as a
                            spreadsheet. One row, several files (D-29).

    Documents generated before S25, or by the Tier 2/3 LLM path, pass none of
    these and carry NULL. That is the honest record, not a gap to backfill.
    """
    import re
    from datetime import datetime
    from database import upload_file, update_document_paths

    record = {
        "user_id": user_id,
        "client_id": client_id,
        "document_type": document_type,
        "language": language,
        "company_name": company_name,
    }

    # Only send stamping columns when they carry a value. Omitting them lets
    # the column defaults apply, and keeps this working if the S25 migration
    # has not been applied to a given environment yet.
    if template_version_id:
        record["template_version_id"] = template_version_id
    if document_group_id:
        record["document_group_id"] = document_group_id
    if outstanding_fields is not None:
        record["outstanding_fields"] = outstanding_fields
    if jurisdictions_applied:
        record["jurisdictions_applied"] = jurisdictions_applied
    if brand_profile_version:
        record["brand_profile_version"] = brand_profile_version

    try:
        supabase = get_supabase()
        res = supabase.table("documents").insert(record).execute()
        doc_id = res.data[0]["id"] if res.data else None
    except Exception as e:
        st.error(f"Could not save document record: {e}")
        return None

    if not doc_id:
        return None

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_company = re.sub(r"[^a-zA-Z0-9_-]", "_", company_name)[:30]
    # RECOSA_, not COMPLAI_. Existing objects keep their old paths — the
    # stored path is what the row points at, so renaming here affects new
    # uploads only and breaks nothing already saved.
    base_path = f"{user_id}/{client_id or 'advisory'}/RECOSA_{document_type}_{safe_company}_{timestamp}"

    docx_path = upload_file(
        "compliance-files", f"{base_path}.docx", docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    pdf_path = None
    if pdf_bytes:
        pdf_path = upload_file(
            "compliance-files", f"{base_path}.pdf", pdf_bytes, "application/pdf"
        )

    # Upload ODT if available
    odt_path = None
    if odt_bytes:
        odt_path = upload_file(
            "compliance-files", f"{base_path}.odt", odt_bytes,
            "application/vnd.oasis.opendocument.text"
        )

    xlsx_path = None
    if xlsx_bytes:
        xlsx_path = upload_file(
            "compliance-files", f"{base_path}.xlsx", xlsx_bytes,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    update_document_paths(doc_id, user_id, docx_path, pdf_path, odt_path,
                          file_path_xlsx=xlsx_path)

    # Auto-register in client document repository as current version
    if client_id and docx_path:
        from database import register_client_document
        register_client_document(
            user_id=user_id,
            client_id=client_id,
            document_type=document_type,
            file_path=docx_path,
            # "recosa_generated", not "complai_generated". Rows written
            # before S26 keep the old string; database.document_source_label()
            # resolves both, and they are NOT backfilled — the old string is
            # the honest record of what wrote them.
            source="recosa_generated",
            change_comment=f"Generated by RECOSA on {datetime.now().strftime('%Y-%m-%d')}",
        )

    # Audit trail (S21) — only logged when generated for a specific client,
    # since audit_log.company_id is NOT NULL and Advisory flows may not have
    # a client_id selected yet
    if client_id:
        from database import log_audit_event
        log_audit_event(
            company_id=client_id,
            user_id=user_id,
            event_type="document_generated",
            event_subtype=document_type,
            resource_id=doc_id,
            summary=f"Generated {DOCUMENT_TYPES.get(document_type, document_type)} document",
        )

    return doc_id



def suggest_processing_activities(client: dict, user_id: str | None = None, client_id: str | None = None) -> dict:
    """
    Use Mistral to suggest likely processing activities, third-party processors
    and retention periods based on the client profile.
    Returns a dict with keys: activities, processors, retention.
    """
    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        raise ValueError("MISTRAL_API_KEY not set")

    sector = client.get("sector", "Unknown")
    size = client.get("company_size", "Unknown")
    country = client.get("country", "BE")
    regulations = client.get("regulations", ["GDPR"])
    company_name = client.get("company_name", "the company")

    country_names = {"BE": "Belgium", "FR": "France", "NL": "Netherlands",
                     "DE": "Germany", "LU": "Luxembourg"}
    country_name = country_names.get(country, country)

    reg_str = ", ".join(regulations) if isinstance(regulations, list) else str(regulations)

    system_prompt = """You are a GDPR compliance expert helping SMEs identify their personal data processing activities.
Based on the company profile provided, generate a realistic and comprehensive list of:
1. Processing activities (what personal data they likely collect and process)
2. Third-party processors (tools and services they likely use)
3. Retention periods (how long they should keep each data type)

IMPORTANT:
- Be realistic and sector-specific — think about what a real company in this sector actually does
- Cover the obvious activities but also less obvious ones (employee data, security logs, etc.)
- For legal basis, choose the most appropriate GDPR Article 6 basis
- For processors, include common tools used in this sector
- Retention periods should follow Belgian/French legal requirements where applicable
- Return ONLY valid JSON, no other text

Return exactly this JSON structure:
{
  "activities": [
    {
      "name": "activity name",
      "subjects": "who the data is about",
      "data": "what personal data",
      "purpose": "why you collect it",
      "legal_basis": "one of: Contract performance (Art. 6(1)(b)) / Consent (Art. 6(1)(a)) / Legal obligation (Art. 6(1)(c)) / Legitimate interests (Art. 6(1)(f)) / Vital interests (Art. 6(1)(d)) / Public task (Art. 6(1)(e))"
    }
  ],
  "processors": [
    {
      "name": "service name",
      "country": "country of the service",
      "purpose": "what it does",
      "data": "what data is shared"
    }
  ],
  "retention": [
    {
      "data_type": "type of data",
      "duration": "how long to keep it"
    }
  ]
}"""

    user_prompt = f"""Company profile:
- Name: {company_name}
- Sector: {sector}
- Size: {size} employees
- Country: {country_name}
- Applicable regulations: {reg_str}

Generate a realistic list of GDPR processing activities for this company.
Include 5-8 processing activities, 3-6 processors, and 4-6 retention rules.
Cover: customer/client data, employee data, marketing, website analytics, and any sector-specific processing."""

    response = requests.post(
        "https://api.mistral.ai/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": "mistral-large-latest",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "max_tokens": 2048,
            "temperature": 0.3,
        },
        timeout=60,
    )
    response.raise_for_status()
    _resp1 = response.json()
    _usage1 = _resp1.get("usage", {})
    try:
        from database import log_token_usage as _ltu
        _ltu(user_id=user_id, feature="docgen_suggest", client_id=client_id,
             input_tokens=_usage1.get("prompt_tokens", 0),
             output_tokens=_usage1.get("completion_tokens", 0))
    except Exception:
        pass

    raw = _resp1["choices"][0]["message"]["content"]

    # Strip markdown fences if present
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip().rstrip("```").strip()

    result = json.loads(raw)

    # Validate structure
    if "activities" not in result:
        result["activities"] = []
    if "processors" not in result:
        result["processors"] = []
    if "retention" not in result:
        result["retention"] = []

    return result

# ── RAG retrieval for document generation ────────────────────

def get_regulatory_context(document_type: str, language: str, country: str) -> str:
    """Retrieve relevant regulatory chunks for document generation."""
    query_map = {
        "privacy_policy": "GDPR privacy policy data controller obligations Articles 13 14 rights",
        "cookie_policy": "ePrivacy cookie consent tracking obligations",
        "dpa": "GDPR data processing agreement Article 28 processor obligations",
        "ropa": "GDPR record of processing activities Article 30",
        "incident_response": "NIS2 incident response plan cybersecurity measures Article 21 23",
        "ai_transparency": "EU AI Act transparency obligations Article 50 AI system disclosure",
    }
    query = query_map.get(document_type, "compliance obligations")

    try:
        chunks = retrieve(
            query=query,
            chunks=[],
            embeddings=None,
            top_k=12,
            language=language,
            country=country,
        )
        parts = [f"[{c.source}]\n{c.text}" for c in chunks]
        return "\n\n---\n\n".join(parts)
    except Exception:
        return ""


# ── Mistral document generation ───────────────────────────────

def generate_document_text(
    document_type: str,
    intake: dict,
    client: dict,
    language: str,
    regulatory_context: str,
    user_id: str | None = None,
    client_id: str | None = None,
) -> str:
    """Call Mistral to generate the document text."""
    api_key = os.environ.get("MISTRAL_API_KEY")
    if not api_key:
        raise ValueError("MISTRAL_API_KEY not set")

    today_str = date.today().strftime("%d %B %Y")
    country = intake.get("country") or client.get("country", "BE")
    dpa_contact = DPA_CONTACTS.get(country, DPA_CONTACTS["EU"])

    lang_instructions = {
        "en": "Write in clear, professional British English.",
        "fr": "Rédigez en français juridique professionnel et clair.",
        "nl": "Schrijf in helder, professioneel juridisch Nederlands.",
    }
    lang_instr = lang_instructions.get(language, lang_instructions["en"])

    doc_prompts = {
        "privacy_policy": _privacy_policy_prompt(intake, client, today_str, dpa_contact),
        "cookie_policy": _cookie_policy_prompt(intake, client, today_str),
        "dpa": _dpa_prompt(intake, client, today_str),
        "ropa": _ropa_prompt(intake, client, today_str),
        "incident_response": _incident_response_prompt(intake, client, today_str),
        "ai_transparency": _ai_transparency_prompt(intake, client, today_str),
    }
    doc_prompt = doc_prompts.get(document_type, "")

    system_prompt = f"""You are an expert EU compliance lawyer drafting professional legal documents for SMEs.
{lang_instr}

IMPORTANT RULES:
- Write complete, legally sound documents ready for immediate use
- Use numbered section headings: "1. Title", "2. Title", "1.1 Sub-section" etc.
- Write in flowing legal prose paragraphs
- Use "- item" for bullet lists where needed
- Do NOT use markdown ## headers — use numbered headings only
- Do NOT use **bold** markers — write plain text
- Do NOT use [text](url) markdown links — write plain email addresses
- Do not include placeholder text like [INSERT NAME] — use the actual data provided
- Base all legal references on the regulatory context provided
- Include specific article references where relevant
- Today's date: {today_str}
- Data Protection Authority for this company: {dpa_contact}

REGULATORY CONTEXT:
{regulatory_context}
"""

    response = requests.post(
        "https://api.mistral.ai/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json={
            "model": "mistral-large-latest",
            "temperature": 0.3,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": doc_prompt},
            ],
            "max_tokens": 8000,
        },
        timeout=120,
    )
    response.raise_for_status()
    result = response.json()
    _usage2 = result.get("usage", {})
    try:
        from database import log_token_usage as _ltu
        _ltu(user_id=user_id, feature="docgen", client_id=client_id,
             input_tokens=_usage2.get("prompt_tokens", 0),
             output_tokens=_usage2.get("completion_tokens", 0))
    except Exception:
        pass
    content_text = result["choices"][0]["message"]["content"]
    finish_reason = result["choices"][0].get("finish_reason", "")

    # If truncated, continue generation
    if finish_reason == "length":
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": doc_prompt},
            {"role": "assistant", "content": content_text},
            {"role": "user", "content": "Continue the document from where you left off. Do not repeat what you already wrote. Continue with the next section."},
        ]
        cont_response = requests.post(
            "https://api.mistral.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "mistral-large-latest",
                "temperature": 0.3,
                "messages": messages,
                "max_tokens": 4000,
            },
            timeout=120,
        )
        if cont_response.status_code == 200:
            _cont_data = cont_response.json()
            _usage3 = _cont_data.get("usage", {})
            try:
                from database import log_token_usage as _ltu
                _ltu(user_id=user_id, feature="docgen", client_id=client_id,
                     input_tokens=_usage3.get("prompt_tokens", 0),
                     output_tokens=_usage3.get("completion_tokens", 0))
            except Exception:
                pass
            content_text += "\n" + _cont_data["choices"][0]["message"]["content"]

    return content_text


def _privacy_policy_prompt(intake: dict, client: dict, today: str, dpa: str) -> str:
    company = f"{intake.get('legal_name') or client.get('company_name', 'The Company')} {intake.get('legal_form', '')}".strip()
    country = intake.get('country') or client.get('country', 'BE')
    return f"""Draft a complete GDPR-compliant Privacy Policy for the following company:

COMPANY DETAILS:
- Legal name: {company}
- Country: {country}
- Sector: {client.get('sector', 'Not specified')}
- Website: {intake.get('website_url') or client.get('website_url', 'Not specified')}
- DPO: {intake.get('dpo_name', 'None appointed')} — {intake.get('dpo_email', '')}
- Contact for data requests: {intake.get('contact_email', '')}

PROCESSING ACTIVITIES:
{intake.get('processing_activities', 'Not specified')}

THIRD-PARTY PROCESSORS:
{intake.get('third_party_processors', 'None specified')}

INTERNATIONAL TRANSFERS: {'Yes' if intake.get('international_transfers') else 'No'}

RETENTION PERIODS:
{intake.get('retention_periods', 'Not specified')}

The document MUST explicitly cover ALL of the following sections:

1. Identity and contact details of the data controller
2. DPO contact details (if applicable) or statement that no DPO is required
3. What personal data we collect, from whom, and why
4. Legal basis for EACH processing activity (Article 6 GDPR) — state the specific basis
5. Special category data: if any special category data (health, biometric, racial, religious, etc.)
   is processed, explicitly state the Article 9(2) GDPR legal ground and additional safeguards
6. Data minimisation: for each processing activity, explain why each data category is necessary
   and limited to what is strictly required for the stated purpose (Art. 5(1)(c))
7. Retention periods for each data category — specific durations, not vague statements
8. Recipients and third-party processors with their role and country
9. International transfers: if applicable, state the transfer mechanism (adequacy decision, SCCs, etc.)
10. Data subject rights (Articles 15-22 GDPR) — list ALL rights with concrete procedures:
    - Right of access (Art. 15): how to submit a request, response timeframe (1 month, extendable to 3)
    - Right to rectification (Art. 16): process for correcting inaccurate data
    - Right to erasure (Art. 17): conditions and process for deletion requests
    - Right to restriction (Art. 18): when and how processing can be restricted
    - Right to data portability (Art. 20): format and delivery method
    - Right to object (Art. 21): how to object to processing
    - Rights re automated decision-making (Art. 22): if applicable
    Include: exact contact email/address for requests, 1-month response timeframe, right to appeal
11. How to exercise rights — step-by-step: "Submit your request to [contact], we will respond within
    30 days, you may escalate to {dpa} if unsatisfied"
12. Right to lodge a complaint with {dpa} — include the authority name and website
13. Technical and organisational security measures (Art. 32) — explicitly describe:
    - Encryption (data in transit via TLS, data at rest)
    - Access controls (role-based access, authentication)
    - Pseudonymisation where applied
    - Staff training on data protection
    - Regular security assessments
    - Data breach detection and response procedures
14. Automated decision-making and profiling (if applicable)
15. Changes to this policy — how users will be notified
16. Effective date: {today}"""


def _cookie_policy_prompt(intake: dict, client: dict, today: str) -> str:
    company = f"{intake.get('legal_name') or client.get('company_name', 'The Company')} {intake.get('legal_form', '')}".strip()
    return f"""Draft a complete ePrivacy-compliant Cookie Policy for:

COMPANY: {company}
WEBSITE: {intake.get('website_url') or client.get('website_url', 'Not specified')}
COOKIES USED: {intake.get('third_party_processors', 'Not specified')}

Cover: what cookies are, categories (strictly necessary / analytics / marketing),
specific cookies used with purpose and duration, how to manage/refuse consent,
contact details, effective date: {today}"""


def _dpa_prompt(intake: dict, client: dict, today: str) -> str:
    company = f"{intake.get('legal_name') or client.get('company_name', 'The Company')} {intake.get('legal_form', '')}".strip()
    return f"""Draft a GDPR Article 28 compliant Data Processing Agreement between:

CONTROLLER: {company}
PROCESSOR: {intake.get('processor_name', 'The Processor')} ({intake.get('processor_country', '')})
PURPOSE OF PROCESSING: {intake.get('processing_purpose', 'Not specified')}
PERSONAL DATA INVOLVED: {intake.get('processing_activities', 'Not specified')}

Cover all mandatory Article 28(3) clauses: processing only on instructions,
confidentiality, security measures, sub-processors, data subject rights assistance,
deletion/return of data, audit rights. Effective date: {today}"""


def _ropa_prompt(intake: dict, client: dict, today: str) -> str:
    company = f"{intake.get('legal_name') or client.get('company_name', 'The Company')} {intake.get('legal_form', '')}".strip()
    return f"""Draft a GDPR Article 30 Record of Processing Activities (RoPA) for:

COMPANY: {company}
SECTOR: {client.get('sector', 'Not specified')}
PROCESSING ACTIVITIES: {intake.get('processing_activities', 'Not specified')}
THIRD PARTIES: {intake.get('third_party_processors', 'None')}
RETENTION: {intake.get('retention_periods', 'Not specified')}
DPO: {intake.get('dpo_name', 'None')} — {intake.get('dpo_email', '')}

Format as a structured table for each processing activity covering:
name of activity, purpose, legal basis, categories of data subjects,
categories of data, recipients, transfers, retention period, security measures.
Date: {today}"""


def _incident_response_prompt(intake: dict, client: dict, today: str) -> str:
    company = f"{intake.get('legal_name') or client.get('company_name', 'The Company')} {intake.get('legal_form', '')}".strip()
    return f"""Draft a NIS2-compliant Incident Response Plan for:

COMPANY: {company}
SECTOR: {client.get('sector', 'Not specified')}
COUNTRY: {client.get('country', 'BE')}
INCIDENT CONTACT: {intake.get('incident_response_contact', 'Not specified')}
ESCALATION PROCEDURE: {intake.get('escalation_procedure', 'Not specified')}
SIZE: {client.get('company_size', 'Not specified')} employees

Cover: scope and objectives, incident classification (low/medium/high/critical),
detection and reporting procedures, NIS2 72-hour notification requirement to national
authority (CCB for Belgium / ANSSI for France), internal escalation chain,
containment and recovery procedures, post-incident review, training requirements.
Date: {today}"""


def _ai_transparency_prompt(intake: dict, client: dict, today: str) -> str:
    company = f"{intake.get('legal_name') or client.get('company_name', 'The Company')} {intake.get('legal_form', '')}".strip()
    return f"""Draft an EU AI Act Article 50 compliant AI System Transparency Notice for:

COMPANY: {company}
AI SYSTEM DESCRIPTION: {intake.get('processing_activities', 'AI-powered system interacting with users')}
CONTACT: {intake.get('contact_email', '')}

Cover: disclosure that users are interacting with an AI system, purpose of the AI system,
limitations and when human oversight applies, how to request human review,
data processed by the AI system, contact for questions. Date: {today}"""


# ── DOCX builder ──────────────────────────────────────────────

# Document footer, per language. Kept beside build_docx rather than in the
# template bodies because it is applied to LLM-generated documents too, which
# have no template to carry it.
_FOOTER = {
    "en": "Generated by RECOSA \u00b7 recosa.eu \u00b7 {today} \u00b7 "
          "This document is a starting point and should be reviewed by a "
          "qualified legal professional.",
    "fr": "G\u00e9n\u00e9r\u00e9 par RECOSA \u00b7 recosa.eu \u00b7 {today} \u00b7 "
          "Ce document constitue un point de d\u00e9part et doit \u00eatre revu par "
          "un professionnel du droit qualifi\u00e9.",
    "nl": "Gegenereerd door RECOSA \u00b7 recosa.eu \u00b7 {today} \u00b7 "
          "Dit document is een uitgangspunt en dient te worden nagekeken door "
          "een gekwalificeerde juridische professional.",
    "de": "Erstellt von RECOSA \u00b7 recosa.eu \u00b7 {today} \u00b7 "
          "Dieses Dokument ist ein Ausgangspunkt und sollte von einer "
          "qualifizierten Rechtsberatung gepr\u00fcft werden.",
}


def build_docx(document_text: str, document_type: str,
               company_name: str, language: str,
               include_header: bool = True) -> bytes:
    """Build a professional DOCX from markdown.

    ── S25 rewrite ───────────────────────────────────────────────────────
    Four defects fixed, three of them silently corrupting every document
    generated before now:

    1. sanitize() ran `re.sub(r'[--]', '', text)`. That was a
       control-character class whose literal control bytes were stripped from
       the source at some point, leaving a range from '-' to '-' — a literal
       hyphen. EVERY hyphen was being deleted from EVERY document:
       "Qu'est-ce" became "Qu'estce", "sous-traitant" became "soustraitant".
       Same class of accident as the `log` find-and-replace incident.

    2. parse_inline() ran `re.sub(r'\\[([^\\]]+)\\]\\([^\\)]+\\)', r'', text)`.
       The replacement was `r'\\1'` — the backslash-one is gone, so markdown
       links were DELETED entirely rather than reduced to their text. A
       document referring to "[the supervisory authority](https://...)" lost
       the phrase as well as the URL.

    3. Markdown tables were not parsed at all. The line simply fell through to
       body text and, because a pipe row looks like nothing else, produced a
       row of pipes — or vanished. This is how the first templated Cookie
       Policy shipped with its third-party vendor disclosure missing while
       reading as a complete document.

    4. The parser was line-oriented, so a hard-wrapped source produced one
       paragraph per LINE rather than per paragraph. Consecutive body lines
       are now joined, with two trailing spaces honoured as a markdown hard
       break (used by multi-line address blocks).

    Also: sanitize() was defined twice and applied twice; em dashes are kept
    rather than degraded to "--"; the footer says RECOSA.

    include_header=False suppresses the generated title block for documents
    whose body already carries its own — otherwise a templated document shows
    its title, company and date twice.
    """
    import re

    # Control characters that lxml cannot serialise. Tab, newline and carriage
    # return are deliberately absent from this class: they are legal in XML and
    # removing them would destroy the document's structure.
    _CONTROL_CHARS = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')

    def sanitize(text: str) -> str:
        if not text:
            return ""
        text = _CONTROL_CHARS.sub('', text)
        text = text.replace('\u00a0', ' ')          # non-breaking space
        text = text.replace('\u2019', "'").replace('\u2018', "'")
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2026', '...')
        # En and em dashes are valid XML and read correctly in Word. The old
        # code degraded '—' to '--', which looks like a typo in a legal text.
        return text.encode('utf-8', errors='ignore').decode('utf-8')

    document_text = sanitize(document_text)
    company_name = sanitize(company_name)

    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc_title = DOCUMENT_TYPES.get(document_type, "Compliance Document")

    # Date in the document's language. strftime('%B') is English regardless of
    # locale unless the locale is installed, which is not guaranteed on
    # Streamlit Cloud — and setlocale() is process-global and not thread-safe.
    try:
        from template_store import format_date
        today = format_date(date.today(), language)
    except Exception:
        today = date.today().strftime("%d %B %Y")

    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    def set_run_font(run, size=10, bold=False, italic=False,
                     color=(0x33, 0x33, 0x33)):
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(*color)

    def add_para_spacing(p, before=0, after=6):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    def add_border_bottom(p, color="D3D1C7"):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def parse_inline(p, text):
        """Inline markdown: ***b+i***, **bold**, *italic*, `code`, [text](url).

        A literal "\\n" in the joined text is a markdown hard break and becomes
        a line break inside the paragraph, not a new paragraph.
        """
        if not text:
            return
        text = _CONTROL_CHARS.sub('', text)
        # Links reduce to their TEXT. The old code substituted an empty string,
        # deleting the phrase along with the URL.
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

        for line_idx, segment in enumerate(text.split('\n')):
            if line_idx:
                p.add_run().add_break()
            parts = re.split(
                r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', segment
            )
            for part in parts:
                if not part:
                    continue
                if part.startswith('***') and part.endswith('***'):
                    set_run_font(p.add_run(part[3:-3]), bold=True, italic=True)
                elif part.startswith('**') and part.endswith('**'):
                    set_run_font(p.add_run(part[2:-2]), bold=True)
                elif part.startswith('*') and part.endswith('*'):
                    set_run_font(p.add_run(part[1:-1]), italic=True)
                elif part.startswith('`') and part.endswith('`'):
                    set_run_font(p.add_run(part[1:-1]), color=(0x44, 0x44, 0x44))
                else:
                    set_run_font(p.add_run(part))

    # ── Markdown table support ────────────────────────────────
    _TABLE_ROW = re.compile(r'^\s*\|.*\|\s*$')
    _TABLE_SEP = re.compile(r'^\s*\|[\s:|-]+\|\s*$')

    def split_row(line):
        cells = line.strip().strip('|').split('|')
        # Restore pipes escaped by the block renderer so a vendor name
        # containing '|' does not silently split the row.
        return [c.strip().replace('\\|', '|') for c in cells]

    def add_table(rows):
        header, body = rows[0], rows[1:]
        table = doc.add_table(rows=1, cols=len(header))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for idx, cell_text in enumerate(header):
            cell = table.rows[0].cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            add_para_spacing(p, before=2, after=2)
            run = p.add_run(cell_text)
            set_run_font(run, size=9, bold=True, color=(0x1B, 0x2A, 0x4A))

        for row in body:
            cells = table.add_row().cells
            for idx in range(len(header)):
                value = row[idx] if idx < len(row) else ""
                cells[idx].text = ""
                p = cells[idx].paragraphs[0]
                add_para_spacing(p, before=2, after=2)
                parse_inline(p, value)
                for run in p.runs:
                    run.font.size = Pt(9)

        doc.add_paragraph()

    # ── Document header ───────────────────────────────────────
    if include_header:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_para_spacing(p, before=0, after=4)
        run = p.add_run(doc_title)
        run.font.name = "Arial"
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        add_border_bottom(p, "1B2A4A")

        p2 = doc.add_paragraph()
        add_para_spacing(p2, before=6, after=2)
        run2 = p2.add_run(company_name)
        run2.font.name = "Arial"
        run2.font.size = Pt(13)
        run2.font.color.rgb = RGBColor(0x4A, 0x3B, 0x8C)

        p3 = doc.add_paragraph()
        add_para_spacing(p3, before=0, after=16)
        run3 = p3.add_run(today)
        run3.font.name = "Arial"
        run3.font.size = Pt(10)
        run3.font.italic = True
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Parse document text ───────────────────────────────────
    lines = document_text.strip().split('\n')
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Skip duplicate title lines that the LLM path often adds.
        if line.strip('# ').strip() in (doc_title, company_name):
            i += 1
            continue

        if not line.strip():
            p = doc.add_paragraph()
            add_para_spacing(p, before=0, after=2)
            i += 1
            continue

        # Table: a pipe row followed by a separator row.
        if (_TABLE_ROW.match(line) and i + 1 < len(lines)
                and _TABLE_SEP.match(lines[i + 1])):
            rows = [split_row(line)]
            j = i + 2
            while j < len(lines) and _TABLE_ROW.match(lines[j].rstrip()):
                rows.append(split_row(lines[j]))
                j += 1
            add_table(rows)
            i = j
            continue

        # H1: '# Title' / '## Title' / '1. Title'
        if re.match(r'^#{1,2}\s+', line) or re.match(r'^\d+\.\s+[A-Z]', line):
            m = re.match(r'^(\d+\.\s+)(.*)', line)
            if m:
                clean = m.group(1) + re.sub(r'^#{1,2}\s+', '', m.group(2))
            else:
                clean = re.sub(r'^#{1,2}\s+', '', line)
            # Word's Heading style, THEN the direct formatting.
            #
            # Direct run formatting alone renders correctly but leaves every
            # paragraph semantically Normal: the navigation pane is empty, a
            # table of contents cannot be built, and assistive software sees
            # flat text. For a record that gets forwarded, opened by someone
            # who did not make it, and read alongside thirty others, structure
            # is not decoration.
            #
            # The style is applied first so the explicit run properties below
            # override its font, size and colour — the appearance is unchanged.
            p = doc.add_paragraph()
            try:
                p.style = doc.styles["Heading 1"]
            except KeyError:
                pass  # a template without the style still renders, unstyled
            add_para_spacing(p, before=14, after=4)
            add_border_bottom(p, "D3D1C7")
            run = p.add_run(clean.strip())
            run.font.name = "Arial"
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
            i += 1
            continue

        # H2
        if re.match(r'^###\s+', line) or re.match(r'^\d+\.\d+\s+[A-Z]', line):
            clean = re.sub(r'^#{2,3}\s+', '', line).strip()
            p = doc.add_paragraph()
            try:
                p.style = doc.styles["Heading 2"]
            except KeyError:
                pass
            add_para_spacing(p, before=10, after=3)
            run = p.add_run(clean)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x4A, 0x3B, 0x8C)
            i += 1
            continue

        # Horizontal rule — checked BEFORE bullets, since '---' also matches
        # the bullet pattern's leading '-'.
        if re.match(r'^-{3,}$', line.strip()):
            p = doc.add_paragraph()
            add_border_bottom(p, "D3D1C7")
            add_para_spacing(p, before=4, after=4)
            i += 1
            continue

        # Bullet
        if re.match(r'^[-*\u2022]\s+', line):
            clean = re.sub(r'^[-*\u2022]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_para_spacing(p, before=1, after=1)
            p.paragraph_format.left_indent = Cm(1)
            parse_inline(p, clean)
            i += 1
            continue

        # Numbered list (not a section heading)
        m_num = re.match(r'^(\d+)\.\s+(.+)', line)
        if m_num and not re.match(r'^\d+\.\s+[A-Z][A-Z]', line):
            p = doc.add_paragraph(style='List Number')
            add_para_spacing(p, before=1, after=1)
            parse_inline(p, m_num.group(2))
            i += 1
            continue

        # Body paragraph — join consecutive lines.
        #
        # Source markdown is hard-wrapped, so one line is NOT one paragraph.
        # Two or more trailing spaces are a markdown hard break and become a
        # line break within the paragraph; anything else joins with a space.
        # This is what keeps a three-line address block intact while stopping
        # every wrapped sentence from becoming its own paragraph.
        chunk = []
        while i < len(lines):
            candidate = lines[i]
            stripped = candidate.rstrip()
            if not stripped.strip():
                break
            if (re.match(r'^#{1,3}\s+', stripped)
                    or re.match(r'^[-*\u2022]\s+', stripped)
                    or re.match(r'^\d+\.\s+', stripped)
                    or re.match(r'^-{3,}$', stripped.strip())
                    or _TABLE_ROW.match(stripped)):
                break
            hard_break = len(candidate) - len(stripped) >= 2
            chunk.append(stripped + ('\n' if hard_break else ''))
            i += 1

        text = ""
        for part in chunk:
            if not text:
                text = part
            elif text.endswith('\n'):
                text += part
            else:
                text += ' ' + part

        p = doc.add_paragraph()
        add_para_spacing(p, before=0, after=5)
        p.paragraph_format.line_spacing = Pt(14)
        parse_inline(p, text)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_border_bottom(p, "D3D1C7")
    add_para_spacing(p, before=0, after=4)
    p = doc.add_paragraph()
    add_para_spacing(p, before=4, after=0)
    # The footer follows the DOCUMENT's language, not the interface's.
    #
    # It was English regardless, so a French register closed with an English
    # sentence — the last untranslated line in an otherwise French artefact,
    # and the one a reader's eye lands on last. The date already resolved per
    # language; the sentence around it did not.
    #
    # Falls back to English for a language with no entry, which is visibly odd
    # rather than blank — the same rule label_for() follows for an unknown code.
    run = p.add_run(
        f"{_FOOTER.get(language, _FOOTER['en']).format(today=today)}"
    )
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF using LibreOffice."""
    import shutil

    # Find soffice.py — try multiple locations
    soffice_candidates = [
        '/mnt/skills/public/docx/scripts/office/soffice.py',
        '/mnt/skills/public/pptx/scripts/office/soffice.py',
    ]
    soffice_script = next((p for p in soffice_candidates if os.path.exists(p)), None)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'document.docx')
        pdf_path = os.path.join(tmpdir, 'document.pdf')

        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        if soffice_script:
            result = subprocess.run(
                ['python3', soffice_script,
                 '--headless', '--convert-to', 'pdf', docx_path, '--outdir', tmpdir],
                capture_output=True, text=True, timeout=60
            )
        else:
            # Fallback: call soffice directly
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', tmpdir, docx_path],
                capture_output=True, text=True, timeout=60
            )

        if os.path.exists(pdf_path):
            with open(pdf_path, 'rb') as f:
                return f.read()
        raise RuntimeError(f"PDF conversion failed: {result.stderr or result.stdout}")


def convert_docx_to_odt(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to ODT using LibreOffice."""
    soffice_candidates = [
        '/mnt/skills/public/docx/scripts/office/soffice.py',
        '/mnt/skills/public/pptx/scripts/office/soffice.py',
    ]
    soffice_script = next((p for p in soffice_candidates if os.path.exists(p)), None)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, 'document.docx')
        odt_path = os.path.join(tmpdir, 'document.odt')

        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)

        if soffice_script:
            result = subprocess.run(
                ['python3', soffice_script,
                 '--headless', '--convert-to', 'odt', docx_path, '--outdir', tmpdir],
                capture_output=True, text=True, timeout=60
            )
        else:
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'odt',
                 '--outdir', tmpdir, docx_path],
                capture_output=True, text=True, timeout=60
            )

        if os.path.exists(odt_path):
            with open(odt_path, 'rb') as f:
                return f.read()
        raise RuntimeError(f"ODT conversion failed: {result.stderr or result.stdout}")
