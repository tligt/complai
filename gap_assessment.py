import os
import io
import json
import time
import re
import requests
import streamlit as st
from datetime import date, datetime
from pypdf import PdfReader
from docx import Document as DocxDocument

# ── Obligation registry ───────────────────────────────────────
# Single source of truth now lives in obligations.py. Re-exported here so
# existing importers (pages/gap.py) keep working unchanged.
from obligations import (            # noqa: F401
    DOCUMENT_TYPES, OBLIGATIONS, PROFILE_QUESTIONS,
    DOC_OBLIGATIONS, DOC_SCORING_OBLIGATIONS, OBLIGATION_TO_DOC,
    OBLIGATION_BY_ID, DOCUMENT_OBLIGATIONS, OPERATIONAL_OBLIGATIONS,
    DOC_CATALOG, REG_DOCS, REGULATION_LABELS,
    KIND_DOCUMENT, KIND_OPERATIONAL,
    is_in_force, applies_from_label, split_by_force,
)

# Status vocabulary. NOT_ASSESSED is deliberately distinct from "missing":
# "missing" asserts the client does not meet an obligation we examined,
# NOT_ASSESSED says we had no evidence to examine. Conflating them either
# defames the client or flatters them, and it wrecks the score either way.
COMPLIANT      = "compliant"
PARTIAL        = "partial"
MISSING        = "missing"
NOT_APPLICABLE = "not_applicable"
NOT_ASSESSED   = "not_assessed"

UNSCORED_STATUSES = (NOT_APPLICABLE, NOT_ASSESSED)

# Relative weights, renormalised over whichever regulations actually have
# assessable obligations — so a client who does not use AI is not dragged
# down by an EU AI Act score of zero.
REGULATION_WEIGHTS = {
    "GDPR": 0.40, "NIS2": 0.30, "EU_AI_ACT": 0.20, "EPRIVACY": 0.10,
}


# ── Document text extraction ──────────────────────────────────

def extract_text_from_upload(uploaded_file) -> str:
    """Extract text from an uploaded PDF or DOCX file."""
    name = uploaded_file.name.lower()
    content = uploaded_file.read()
    if name.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(p.extract_text() for p in reader.pages if p.extract_text())
        except Exception:
            return ""
    elif name.endswith(".docx"):
        try:
            doc = DocxDocument(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception:
            return ""
    elif name.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    return ""


def extract_text_from_storage(file_path: str) -> str:
    """Extract text from a file stored in Supabase Storage."""
    from database import get_supabase_admin
    import io
    try:
        admin = get_supabase_admin()
        content = admin.storage.from_("compliance-files").download(file_path)

        # Sniff the format from the bytes, not the filename. Upload paths are
        # built with re.sub(r"[^a-zA-Z0-9_-]", "_", name), which turns
        # "policy.pdf" into "policy_pdf" — so extension matching silently
        # failed and every PDF fell through to .decode(), producing binary
        # soup that the model then reported as "a corrupted document".
        if content[:4] == b"%PDF":
            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(p.extract_text() for p in reader.pages
                               if p.extract_text())
        if content[:4] == b"PK\x03\x04":          # zip container => docx
            doc = DocxDocument(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        # Fall back to filename for anything unsniffable, then to plain text.
        name = file_path.lower()
        if name.endswith(".pdf") or "_pdf" in name:
            try:
                reader = PdfReader(io.BytesIO(content))
                return "\n\n".join(p.extract_text() for p in reader.pages
                                   if p.extract_text())
            except Exception:
                pass
        if name.endswith(".docx") or "_docx" in name:
            try:
                doc = DocxDocument(io.BytesIO(content))
                return "\n\n".join(p.text for p in doc.paragraphs
                                   if p.text.strip())
            except Exception:
                pass
        return content.decode("utf-8", errors="ignore")
    except Exception as e:
        return ""


# ── AI analysis ───────────────────────────────────────────────

def analyse_obligation(obligation: dict, document_text: str,
                        profile_answers: dict, api_key: str) -> dict:
    """Analyse a single obligation. Returns status, explanation, recommendation."""
    ob_id = obligation["id"]

    # Profile-based obligations
    if obligation["profile_question"]:
        q_key = obligation["profile_question"]
        answer = profile_answers.get(q_key, "")
        q_config = PROFILE_QUESTIONS.get(q_key, {})
        compliant = q_config.get("compliant_answers", [])
        partial = q_config.get("partial_answers", [])
        na_answers = q_config.get("na_answers", [])

        # Not applicable
        if answer in na_answers and answer not in compliant:
            return {"id": ob_id, "status": "not_applicable",
                    "explanation": f"Not applicable: {answer}", "recommendation": ""}

        if answer in compliant:
            return {"id": ob_id, "status": "compliant",
                    "explanation": f"Confirmed: {answer}", "recommendation": ""}
        elif answer in partial:
            return {"id": ob_id, "status": "partial",
                    "explanation": f"Partially compliant: {answer}",
                    "recommendation": "Formalise with documented evidence of completion."}
        else:
            return {"id": ob_id, "status": "missing",
                    "explanation": f"Not in place: {answer}",
                    "recommendation": f"Implement: {obligation['title']}"}

    # Document not provided
    if not document_text.strip():
        doc_label = DOCUMENT_TYPES.get(obligation.get("doc_type",""), "document")
        return {"id": ob_id, "status": "missing",
                "explanation": f"No {doc_label} provided — cannot assess this obligation.",
                "recommendation": f"Upload or generate a {doc_label} in RECOSA."}

    # Mistral analysis
    if len(document_text) <= 6000:
        doc_excerpt = document_text
    else:
        first = document_text[:3000]
        mid_s = len(document_text)//2 - 500
        middle = document_text[mid_s:mid_s+1000]
        last = document_text[-1500:]
        sep = "\n...[continued]...\n"
        doc_excerpt = first + sep + middle + sep + last
    system_prompt = """You are an EU compliance expert assessing whether a document satisfies a regulatory obligation.
Analyse the document excerpt and respond ONLY with valid JSON:
{
  "status": "compliant" | "partial" | "missing",
  "explanation": "one sentence assessment",
  "recommendation": "one sentence action (empty if compliant)"
}"""

    user_prompt = f"""OBLIGATION: {obligation['title']}
REGULATION: {obligation['regulation']} {obligation['article']}
REQUIREMENT: {obligation['description']}

DOCUMENT:
{doc_excerpt}"""

    for attempt in range(3):
        try:
            response = requests.post(
                "https://api.mistral.ai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}",
                         "Content-Type": "application/json"},
                json={
                    "model": "mistral-large-latest",
                    "temperature": 0.0,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "max_tokens": 300,
                },
                timeout=30,
            )
            response.raise_for_status()
            _rdata1 = response.json()
            _u1 = _rdata1.get("usage", {})
            try:
                from database import log_token_usage as _ltu
                _ltu(user_id=user_id, feature="gap_single", client_id=client_id,
                     input_tokens=_u1.get("prompt_tokens", 0),
                     output_tokens=_u1.get("completion_tokens", 0))
            except Exception:
                pass
            raw = _rdata1["choices"][0]["message"]["content"].strip()
            # Strip markdown fences
            raw = re.sub(r"```json|```", "", raw).strip()
            # Extract JSON object if embedded in text
            json_match = re.search(r'\{[^{}]+\}', raw, re.DOTALL)
            if json_match:
                raw = json_match.group(0)
            result = json.loads(raw)
            # Validate required fields
            if "status" not in result:
                raise ValueError("Missing status field")
            if result["status"] not in ("compliant","partial","missing"):
                raise ValueError(f"Invalid status: {result['status']}")
            result["id"] = ob_id
            result.setdefault("explanation", "")
            result.setdefault("recommendation", "")
            return result
        except Exception as e:
            if attempt < 2:
                time.sleep(5)

    # Final fallback — flag as needs manual review, not missing
    return {"id": ob_id, "status": "partial",
            "explanation": "Could not assess automatically — manual review recommended.",
            "recommendation": f"Manually verify compliance with: {obligation['title']}"}



# ── Document-specific obligations mapping ─────────────────────


def _assess_obligations_batch(
    obligations: list,
    doc_excerpt: str,
    api_key: str,
) -> list:
    """Assess 2-3 obligations in a single Mistral call. Returns list of results."""
    ob_list = "\n".join(
        f"- ID={ob['id']}: {ob['title']} ({ob['article']}): {ob['description']}"
        for ob in obligations
    )

    system_prompt = """You are an EU compliance expert. Check if the document satisfies each obligation.
Return ONLY a valid JSON array, one object per obligation, no other text:
[{"id":"ob_id","status":"compliant","explanation":"one sentence","recommendation":""},...]
Status must be exactly: compliant, partial, or missing."""

    user_prompt = f"""DOCUMENT:
{doc_excerpt}

CHECK THESE OBLIGATIONS:
{ob_list}"""

    for attempt in range(3):
        try:
            response = requests.post(
                "https://api.mistral.ai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}",
                         "Content-Type": "application/json"},
                json={
                    "model": "mistral-large-latest",
                    "temperature": 0.0,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "max_tokens": 800,
                },
                timeout=45,
            )
            response.raise_for_status()
            _rdata2 = response.json()
            _u2 = _rdata2.get("usage", {})
            try:
                from database import log_token_usage as _ltu
                _ltu(user_id=user_id, feature="gap_full", client_id=client_id,
                     input_tokens=_u2.get("prompt_tokens", 0),
                     output_tokens=_u2.get("completion_tokens", 0))
            except Exception:
                pass
            raw = _rdata2["choices"][0]["message"]["content"].strip()
            # Strip markdown fences
            raw = re.sub(r"```json|```", "", raw).strip()
            # Extract JSON array
            match = re.search(r'\[[\s\S]*\]', raw)
            if match:
                raw = match.group(0)
            parsed = json.loads(raw)
            results = []
            results_map = {r.get("id",""): r for r in parsed if isinstance(r, dict)}
            for ob in obligations:
                r = results_map.get(ob["id"], {})
                status = r.get("status","partial")
                if status not in ("compliant","partial","missing"):
                    status = "partial"
                results.append({
                    "id": ob["id"],
                    "status": status,
                    "explanation": r.get("explanation","Could not assess."),
                    "recommendation": r.get("recommendation",""),
                })
            return results
        except Exception:
            if attempt < 2:
                time.sleep(5)

    # Fallback
    return [{"id": ob["id"], "status": "partial",
             "explanation": "Could not assess automatically — manual review recommended.",
             "recommendation": f"Manually verify: {ob['title']}"} for ob in obligations]


def run_document_review(
    doc_type: str,
    document_text: str,
    profile_answers: dict,
    client: dict,
    user_id: str | None = None,
    client_id: str | None = None,
) -> dict:
    """
    Review a single document against its specific obligations.
    Uses batches of 2 obligations per Mistral call to balance
    context quality vs number of API calls.
    """
    api_key = os.environ.get("MISTRAL_API_KEY", "")
    ob_ids = DOC_OBLIGATIONS.get(doc_type, [])
    obligations = [o for o in OBLIGATIONS if o["id"] in ob_ids]

    if not obligations:
        return {"results": [], "score": 0,
                "doc_type": doc_type, "doc_label": DOCUMENT_TYPES.get(doc_type, doc_type)}

    # Separate profile-based from document-based obligations
    profile_results = []
    doc_obligations = []
    for ob in obligations:
        if ob["profile_question"]:
            result = analyse_obligation(ob, "", profile_answers, api_key)
            profile_results.append(result)
        else:
            doc_obligations.append(ob)

    doc_results = []
    if doc_obligations and document_text.strip():
        # 5-chunk sampling — covers full document evenly
        doc_len = len(document_text)
        if doc_len <= 10000:
            doc_excerpt = document_text
        else:
            sep = "\n...[continued]...\n"
            chunk = doc_len // 5
            doc_excerpt = sep.join(
                document_text[i*chunk:i*chunk+2000] for i in range(5)
            )

        # Process in batches of 2 — avoids lost-in-middle while limiting API calls
        total = len(doc_obligations)
        progress = st.progress(0, text="Reviewing document...")
        for i in range(0, total, 2):
            batch = doc_obligations[i:i+2]
            progress.progress(
                (i + len(batch)) / total,
                text=f"Checking obligations {i+1}-{min(i+2,total)} of {total}..."
            )
            batch_results = _assess_obligations_batch(batch, doc_excerpt, api_key)
            doc_results.extend(batch_results)
        progress.empty()

    elif doc_obligations:
        for ob in doc_obligations:
            doc_results.append({
                "id": ob["id"], "status": "missing",
                "explanation": f"No {DOCUMENT_TYPES.get(doc_type,'document')} provided.",
                "recommendation": f"Upload or generate a {DOCUMENT_TYPES.get(doc_type,'document')} in RECOSA."
            })

    # Combine and reorder
    all_results = {r["id"]: r for r in profile_results + doc_results}
    results = [all_results.get(ob["id"], {"id": ob["id"], "status": "missing",
               "explanation": "", "recommendation": ""}) for ob in obligations]

    # Score only on obligations this document actually SATISFIES. The review
    # set is deliberately wider — it asks a privacy policy about data
    # minimisation (gdpr_14) and security measures (gdpr_20) because that is
    # useful feedback — but those are operational obligations that no
    # document can discharge, so counting them would mark a complete policy
    # as merely partial. Reported, not scored.
    scoring_ids = set(DOC_SCORING_OBLIGATIONS.get(doc_type, []))
    scored = [r for r in results if r["id"] in scoring_ids]

    n_compliant = sum(1 for r in scored if r["status"] == "compliant")
    n_partial   = sum(1 for r in scored if r["status"] == "partial")
    score = int((n_compliant + n_partial * 0.5) / len(scored) * 100) if scored else 0

    return {
        "results": results,
        "score": score,
        "scored_ids": sorted(scoring_ids),
        "doc_type": doc_type,
        "doc_label": DOCUMENT_TYPES.get(doc_type, doc_type),
        "obligations": obligations,
    }


def run_gap_assessment(
    current_docs: dict,       # {doc_type: {file_path, version, ...}} from client_documents
    uploaded_docs: dict,      # {doc_type: text} from session uploads (overrides stored)
    profile_answers: dict,
    client: dict,
    user_id: str | None = None,
    client_id: str | None = None,
) -> dict:
    """
    Run gap assessment. For each obligation:
    - If doc_type has an uploaded version → use that text
    - If doc_type has a stored version → fetch and use that text
    - If neither → missing
    - If profile_question → use profile answers
    """
    api_key = os.environ.get("MISTRAL_API_KEY", "")

    # Build document text cache — uploaded takes priority over stored
    doc_texts = {}

    # Load stored documents
    for doc_type, doc_record in current_docs.items():
        if doc_type not in uploaded_docs:
            text = extract_text_from_storage(doc_record["file_path"])
            if text:
                doc_texts[doc_type] = text

    # Add uploaded documents (override)
    doc_texts.update(uploaded_docs)

    results = []
    total = len(OBLIGATIONS)
    progress = st.progress(0, text="Starting gap assessment...")

    for i, obligation in enumerate(OBLIGATIONS):
        progress.progress(
            (i + 1) / total,
            text=f"Analysing {i+1}/{total}: {obligation['title'][:50]}..."
        )

        # Not yet binding: report the date, do not judge the client on it.
        if not is_in_force(obligation):
            results.append({
                "id": obligation["id"], "status": NOT_ASSESSED,
                "explanation": f"Applies from {applies_from_label(obligation)}. "
                               "Not assessed because it is not yet in force.",
                "recommendation": "",
            })
            continue

        doc_type = obligation.get("doc_type")
        doc_text = doc_texts.get(doc_type, "") if doc_type else ""

        # No document to read means we have no evidence either way. Saying
        # "missing" would assert the client fails an obligation we never
        # examined — and it used to drag the score down, because the exclusion
        # filter matched the substring "not provided" against an explanation
        # that actually read "No X provided". Set the status directly instead
        # of inferring it from prose, and skip the API call entirely.
        if doc_type and not doc_text.strip():
            label = DOCUMENT_TYPES.get(doc_type, "document")
            results.append({
                "id": obligation["id"], "status": NOT_ASSESSED,
                "explanation": f"No {label} in the repository — this obligation "
                               "could not be assessed.",
                "recommendation": f"Upload or generate a {label} in RECOSA.",
            })
            continue

        # For obligations with no specific doc_type, combine all available texts
        if not doc_type and not obligation["profile_question"]:
            doc_text = "\n\n---\n\n".join(doc_texts.values())[:4000]

        result = analyse_obligation(obligation, doc_text, profile_answers, api_key)
        results.append(result)

        if obligation["profile_question"] is None:
            time.sleep(0.3)

    progress.empty()

    # ── Scores ────────────────────────────────────────────────
    def calc_score(reg: str):
        """Score for one regulation, or None if nothing was assessable.

        None and 0 mean very different things: None is "we have no basis to
        score this", 0 is "we assessed it and you meet none of it". Returning
        0 for both is how an unassessed regulation used to drag the overall
        score down.
        """
        reg_results = [
            r for r, o in zip(results, OBLIGATIONS)
            if o["regulation"] == reg and r["status"] not in UNSCORED_STATUSES
        ]
        if not reg_results:
            return None
        compliant = sum(1 for r in reg_results if r["status"] == COMPLIANT)
        partial   = sum(1 for r in reg_results if r["status"] == PARTIAL)
        return int((compliant + partial * 0.5) / len(reg_results) * 100)

    raw = {reg: calc_score(reg) for reg in REGULATION_WEIGHTS}
    scored = {reg: s for reg, s in raw.items() if s is not None}

    if scored:
        total_weight = sum(REGULATION_WEIGHTS[r] for r in scored)
        score_overall = int(round(sum(
            s * REGULATION_WEIGHTS[r] for r, s in scored.items()
        ) / total_weight))
    else:
        score_overall = 0

    n_assessed = sum(1 for r in results if r["status"] not in UNSCORED_STATUSES)

    return {
        "results": results,
        # 0 for storage compatibility; scored_regulations says which are real.
        "score_gdpr":      raw["GDPR"]      or 0,
        "score_nis2":      raw["NIS2"]      or 0,
        "score_eprivacy":  raw["EPRIVACY"]  or 0,
        "score_eu_ai_act": raw["EU_AI_ACT"] or 0,
        "score_overall":   score_overall,
        "scored_regulations": sorted(scored),
        "n_assessed": n_assessed,
        "n_total": len(results),
    }


# ── PDF report ────────────────────────────────────────────────

def generate_gap_report_pdf(assessment: dict, client: dict,
                             profile_answers: dict, doc_versions: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, black, white
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, PageBreak)
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm)

    NAVY   = HexColor("#1B2A4A")
    PURPLE = HexColor("#4A3B8C")
    TEAL   = HexColor("#0F6E56")
    RED    = HexColor("#C0392B")
    AMBER  = HexColor("#E67E22")
    GREEN  = HexColor("#27AE60")
    LGRAY  = HexColor("#F4F3F0")
    MGRAY  = HexColor("#888888")

    styles = getSampleStyleSheet()
    def S(name, **kw): return ParagraphStyle(name, parent=styles["Normal"], **kw)

    S_TITLE = S("t", fontSize=22, textColor=NAVY,   leading=28, spaceAfter=4)
    S_SUB   = S("s", fontSize=12, textColor=PURPLE, leading=16, spaceAfter=4)
    S_DATE  = S("d", fontSize=9,  textColor=MGRAY,  leading=12, spaceAfter=16)
    S_H1    = S("h1",fontSize=13, textColor=NAVY,   leading=18, spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold")
    S_H2    = S("h2",fontSize=10, textColor=PURPLE, leading=14, spaceBefore=8,  spaceAfter=3, fontName="Helvetica-Bold")
    S_BODY  = S("b", fontSize=9,  textColor=black,  leading=13, spaceAfter=4)
    S_SMALL = S("sm",fontSize=8,  textColor=MGRAY,  leading=11, spaceAfter=2)
    S_REC   = S("r", fontSize=9,  textColor=TEAL,   leading=13, spaceAfter=5, leftIndent=12)

    today   = date.today().strftime("%d %B %Y")
    company = client.get("company_name", "")
    results_by_id = {r["id"]: r for r in assessment["results"]}

    story = []

    # Header
    story.append(Paragraph("Gap Assessment Report", S_TITLE))
    story.append(HRFlowable(width="100%", thickness=2, color=NAVY))
    story.append(Spacer(1, 4))
    story.append(Paragraph(company, S_SUB))
    story.append(Paragraph(f"Generated by RECOSA · {today}", S_DATE))

    # Document versions used
    if doc_versions:
        story.append(Paragraph("Documents assessed:", S_H2))
        for dt, info in doc_versions.items():
            label = DOCUMENT_TYPES.get(dt, dt)
            src = "RECOSA generated" if info.get("source") == "complai_generated" else "Client upload"
            story.append(Paragraph(
                f"• {label} — v{info.get('version',1)} ({src})", S_SMALL))
        story.append(Spacer(1, 8))

    # Scores table
    story.append(Paragraph("Executive Summary", S_H1))
    story.append(HRFlowable(width="100%", thickness=1, color=LGRAY))
    story.append(Spacer(1, 6))

    def status_label(s):
        if s >= 75: return "Good"
        if s >= 50: return "Needs work"
        return "Critical gaps"

    score_data = [
        ["Regulation", "Score", "Status"],
        ["GDPR",      f"{assessment['score_gdpr']}/100",     status_label(assessment['score_gdpr'])],
        ["NIS2",      f"{assessment['score_nis2']}/100",     status_label(assessment['score_nis2'])],
        ["ePrivacy",  f"{assessment['score_eprivacy']}/100", status_label(assessment['score_eprivacy'])],
        ["EU AI Act", f"{assessment.get('score_eu_ai_act', 0)}/100",
         status_label(assessment.get('score_eu_ai_act', 0))],
        ["OVERALL",   f"{assessment['score_overall']}/100",  status_label(assessment['score_overall'])],
    ]
    t = Table(score_data, colWidths=[8*cm, 4*cm, 5*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0,0),(-1,0), NAVY),
        ("TEXTCOLOR",     (0,0),(-1,0), white),
        ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
        ("FONTSIZE",      (0,0),(-1,-1), 9),
        ("BACKGROUND",    (0,-1),(-1,-1), LGRAY),
        ("FONTNAME",      (0,-1),(-1,-1), "Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,1),(-1,-2), [white, HexColor("#F9F9F9")]),
        ("GRID",          (0,0),(-1,-1), 0.5, HexColor("#DDDDDD")),
        ("PADDING",       (0,0),(-1,-1), 8),
        ("ALIGN",         (1,0),(1,-1), "CENTER"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    statuses = [r["status"] for r in assessment["results"]]
    n_c = statuses.count("compliant")
    n_p = statuses.count("partial")
    n_m = statuses.count("missing")
    n_na = statuses.count(NOT_APPLICABLE)
    n_ns = statuses.count(NOT_ASSESSED)
    n_total = len([s for s in statuses if s not in UNSCORED_STATUSES])

    story.append(Paragraph(
        f"Out of <b>{n_total}</b> applicable obligations: "
        f"<font color='#27AE60'><b>{n_c} compliant</b></font>, "
        f"<font color='#E67E22'><b>{n_p} partial</b></font>, "
        f"<font color='#C0392B'><b>{n_m} missing</b></font>."
        + (f" {n_na} not applicable." if n_na else "")
        + (f" {n_ns} could not be assessed." if n_ns else ""), S_BODY))

    story.append(PageBreak())

    # Gaps by regulation
    for regulation in ["GDPR", "NIS2", "EPRIVACY", "EU_AI_ACT"]:
        reg_labels = {"GDPR":"GDPR","NIS2":"NIS2 Directive",
                      "EPRIVACY":"ePrivacy Directive","EU_AI_ACT":"EU AI Act"}
        reg_obs = [o for o in OBLIGATIONS if o["regulation"] == regulation]
        has_gaps = any(
            results_by_id.get(o["id"],{}).get("status") in ["partial","missing"]
            for o in reg_obs
        )
        if not has_gaps:
            continue

        story.append(Paragraph(reg_labels[regulation], S_H1))
        story.append(HRFlowable(width="100%", thickness=1, color=LGRAY))

        for priority in ["high","medium","low"]:
            p_obs = [o for o in reg_obs if o["priority"] == priority]
            if not p_obs: continue
            p_gaps = [o for o in p_obs
                      if results_by_id.get(o["id"],{}).get("status") in ["partial","missing"]]
            if not p_gaps: continue

            p_labels = {"high":"🔴 High Priority","medium":"🟡 Medium Priority","low":"🔵 Low Priority"}
            story.append(Paragraph(p_labels[priority], S_H2))

            for ob in p_gaps:
                result = results_by_id.get(ob["id"], {})
                icon = "⚠️" if result.get("status") == "partial" else "❌"
                story.append(Paragraph(
                    f"{icon} <b>{ob['title']}</b> <font color='#888888'>({ob['article']})</font>",
                    S_BODY))
                story.append(Paragraph(result.get("explanation",""), S_SMALL))
                if result.get("recommendation"):
                    story.append(Paragraph(f"→ {result['recommendation']}", S_REC))
                if ob.get("doc_type"):
                    doc_label = DOCUMENT_TYPES.get(ob["doc_type"], ob["doc_type"])
                    story.append(Paragraph(
                        f"📄 Generate a compliant {doc_label} in RECOSA → Documents page", S_REC))

        story.append(Spacer(1, 8))

    story.append(PageBreak())

    # Compliant
    story.append(Paragraph("What's Compliant", S_H1))
    story.append(HRFlowable(width="100%", thickness=1, color=LGRAY))
    story.append(Spacer(1, 4))
    compliant_obs = [o for o in OBLIGATIONS
                     if results_by_id.get(o["id"],{}).get("status") == "compliant"]
    if compliant_obs:
        for ob in compliant_obs:
            story.append(Paragraph(
                f"✅ <b>{ob['title']}</b> ({ob['regulation']} {ob['article']})", S_BODY))
    else:
        story.append(Paragraph("No fully compliant obligations identified.", S_BODY))

    story.append(Spacer(1, 12))

    # Next steps
    story.append(Paragraph("Recommended Next Steps", S_H1))
    story.append(HRFlowable(width="100%", thickness=1, color=LGRAY))
    story.append(Spacer(1, 4))
    step = 1
    for priority in ["high","medium","low"]:
        for ob in OBLIGATIONS:
            if ob["priority"] != priority: continue
            result = results_by_id.get(ob["id"],{})
            if result.get("status") not in ["partial","missing"]: continue
            story.append(Paragraph(
                f"{step}. <b>{ob['title']}</b> ({ob['regulation']} {ob['article']})", S_BODY))
            if result.get("recommendation"):
                story.append(Paragraph(f"   {result['recommendation']}", S_REC))
            step += 1
            if step > 10: break
        if step > 10: break

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=1, color=LGRAY))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        f"Generated by RECOSA · recosa.eu · {today} · "
        "AI analysis of provided documents and self-reported information. "
        "Review with a qualified legal or cybersecurity professional.", S_SMALL))

    doc.build(story)
    return buf.getvalue()


# ── Persistence ───────────────────────────────────────────────

def save_gap_assessment(user_id: str, client_id: str | None,
                         assessment: dict, profile_answers: dict,
                         pdf_bytes: bytes) -> str | None:
    from database import get_supabase, upload_file
    try:
        supabase = get_supabase()
        res = supabase.table("gap_assessments").insert({
            "user_id": user_id,
            "client_id": client_id,
            "regulations": assessment.get("scored_regulations")
                           or ["GDPR","NIS2","EPRIVACY"],
            "score_gdpr": assessment["score_gdpr"],
            "score_nis2": assessment["score_nis2"],
            "score_eprivacy": assessment["score_eprivacy"],
            "score_eu_ai_act": assessment.get("score_eu_ai_act", 0),
            "score_overall": assessment["score_overall"],
            "gaps": assessment["results"],
            "profile_answers": profile_answers,
        }).execute()
        assessment_id = res.data[0]["id"] if res.data else None
    except Exception as e:
        st.warning(f"Could not save assessment: {e}")
        return None

    if pdf_bytes and assessment_id:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = f"{user_id}/{client_id or 'advisory'}/gap_reports/RECOSA_gap_{ts}.pdf"
        stored = upload_file("compliance-files", path, pdf_bytes, "application/pdf")
        if stored:
            try:
                get_supabase().table("gap_assessments") \
                    .update({"file_path_pdf": path}) \
                    .eq("id", assessment_id).execute()
            except Exception:
                pass

    # Audit trail (S21) — only logged when there's a specific client
    if client_id and assessment_id:
        from database import log_audit_event
        log_audit_event(
            company_id=client_id,
            user_id=user_id,
            event_type="gap_assessment_run",
            event_subtype="full",
            resource_id=assessment_id,
            summary="Ran full gap assessment",
            metadata={"score_overall": assessment["score_overall"],
                      "score_gdpr": assessment["score_gdpr"],
                      "score_nis2": assessment["score_nis2"],
                      "score_eprivacy": assessment["score_eprivacy"],
                      "score_eu_ai_act": assessment.get("score_eu_ai_act", 0)},
        )

    return assessment_id


def load_gap_assessment_history(user_id: str, client_id: str | None) -> list[dict]:
    from database import get_supabase
    try:
        supabase = get_supabase()
        q = supabase.table("gap_assessments") \
            .select("id,created_at,score_overall,score_gdpr,score_nis2,score_eprivacy,score_eu_ai_act,file_path_pdf") \
            .eq("user_id", user_id) \
            .order("created_at", desc=True).limit(10)
        if client_id:
            q = q.eq("client_id", client_id)
        return q.execute().data or []
    except Exception:
        return []
